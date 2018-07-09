import unittest
from collections import deque
from util.msword import DocWriter

import random
import string

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os


class TestClass(unittest.TestCase):

    def test_add_paragraph(self):
        self.doc_writer.add_paragraph(self.doc_writer.document, self.heading, self.lines)
        self.assertEqual(len(self.lines)+1, len(self.doc_writer.document.paragraphs))

        text=''
        for p in self.doc_writer.document.paragraphs:
            text+=p.text

        self.assertEqual(self.heading+''.join(self.lines), text)

    def test_add_bold_paragraph(self):
        self.doc_writer.add_paragraph_bold_checking(self.doc_writer.document, self.heading, self.lines, 'あ')
        self.assertEqual(len(self.lines)+1, len(self.doc_writer.document.paragraphs))

        text=''
        for p in self.doc_writer.document.paragraphs:
            text+=p.text

        self.assertEqual(self.heading+''.join(self.lines), text)


    def test_write_to_doc(self):

        f = open('C:\\Users\\01556729\\Dropbox\\Code\\JAccent\\data\\test_data.txt', encoding='utf-8')
        lines = f.readlines()
        f.close()

        heading1='Japanese Accent'
        heading2='Dictionary Note\n'
        heading3='Accent Note\n'
        heading4='No Result Note\n'

        self.doc_writer.write_to_doc(lines, [],[],[])
        text = ''
        for p in self.doc_writer.document.paragraphs:
            text += p.text

        self.assertEqual(len(heading1+heading2+heading3+heading4+''.join(lines)), len(text))

    def setUp(self):
        self.doc_writer = DocWriter('C:\\Users\\01556729\\Dropbox\\Code\\JAccent\\data\\', 'Test.doc')
        #self.heading = "".join([random.choice(string.ascii_letters) for i in range(0,15)])
        self.heading = "Heading"
        self.lines = []
        for i in range(0, 100):
            self.lines.append("".join([random.choice(string.ascii_letters) for i in range(0,15)]))


    if __name__ == '__main__':
        unittest.main()
