from collections import deque

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


class DocWriter:
    def __init__(self, project_data, doc_filename):
        self.project_data = project_data
        self.doc_filename = doc_filename

    def open_new_document(self, doc_filename):
        self.doc_filename = doc_filename
        self.document = Document()

    def add_paragraph(self, document, heading, lines):
        document.add_heading(heading)
        for line in lines:
            paragraph = document.add_paragraph()
            paragraph.add_run(line)

    def add_paragraph_bold_checking(self, document, heading, lines, keyword):
        document.add_heading(heading)
        for line in lines:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            if keyword in line:
                run.font.bold = True

    def get_accent_as_number(self, accent):
        if accent == 'α':
            return '0'
        elif accent == 'β':
            return '1'
        elif accent == 'γ':
            return '2'
        elif accent == 'δ':
            return '3'
        elif accent == 'ϵ':
            return '4'
        elif accent == 'ζ':
            return '5'
        elif accent == 'η':
            return '6'

    def write_to_doc(self, output_line, accent_note, debug_note, no_result_note):
        accent_symbol = ['α', 'β', 'γ', 'δ', 'ϵ', 'ζ', 'η']
        first_no_voice = ['き', 'し', 'ち', 'ひ', 'ぴ', 'く', 'す', 'つ', 'ふ', 'ぷ']
        second_no_voice = ['か', 'き', 'く', 'け', 'こ', 'さ', 'し', 'す', 'せ', 'そ', 'た', 'ち', 'つ', 'て', 'と', 'は', 'ひ', 'ふ',
                           'へ', 'ほ', 'ぱ', 'ぴ', 'ぷ', 'ぺ', 'ぽ']
        newline_symbol = 'Ω'
        quickbeat_letter = 'っ'

        self.document.add_heading('Japanese Accent')
        paragraph = self.document.add_paragraph()

        run_list = []
        for char in output_line:
            if char in accent_symbol:
                accent = self.get_accent_as_number(char)
                run = paragraph.add_run(accent)
                run.font.superscript = True
            elif char == newline_symbol:
                paragraph = self.document.add_paragraph()
                continue
            else:
                run = paragraph.add_run(char)
            run_list.append(run)

        run_stack = deque(run_list)

        while len(run_stack) > 1:
            if run_stack[1]._r.text == quickbeat_letter:
                if run_stack[0]._r.text in first_no_voice and run_stack[2]._r.text in second_no_voice:
                    run_stack[0].font.highlight_color = WD_COLOR_INDEX.GRAY_25
                run_stack.popleft()
                run_stack.popleft()
            elif run_stack[0]._r.text in first_no_voice and run_stack[1]._r.text in second_no_voice:
                run_stack[0].font.highlight_color = WD_COLOR_INDEX.GRAY_25
                run_stack.popleft()
            else:
                run_stack.popleft()

        self.document.add_page_break()
        self.add_paragraph_bold_checking(self.document, 'Dictionary Note', debug_note, '請覆查字典')

        self.document.add_page_break()
        self.add_paragraph(self.document, 'Accent Note', accent_note)

        self.document.add_page_break()
        self.add_paragraph(self.document, 'No Result Note', no_result_note)

        self.document.save(self.project_data + self.doc_filename)

